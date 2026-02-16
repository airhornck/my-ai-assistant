"""
Word 报告生成器
通用报告生成插件，支持生成各类 Word 文档报告
"""
from __future__ import annotations

import logging
import os
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """Word 文档报告生成器"""

    def __init__(self, output_dir: str = "data/reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_report(
        self,
        report_type: str,
        title: str,
        summary: str,
        sections: List[Dict[str, str]],
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        生成 Word 报告文档

        Args:
            report_type: 报告类型 (account_diagnosis, marketing_strategy, viral_prediction, custom)
            title: 报告标题
            summary: 执行摘要
            sections: 报告章节列表 [{"heading": "章节标题", "content": "章节内容"}]
            metadata: 元数据 {"author": "...", "brand": "...", "created_at": "..."}

        Returns:
            生成的文件路径
        """
        doc = Document()

        # 设置文档标题
        self._add_title(doc, title)

        # 添加元数据
        if metadata:
            self._add_metadata(doc, metadata)

        # 添加摘要
        self._add_summary(doc, summary)

        # 添加分隔线
        doc.add_paragraph("_" * 50)

        # 添加各章节
        for section in sections:
            self._add_section(
                doc,
                section.get("heading", ""),
                section.get("content", ""),
            )

        # 添加页脚
        self._add_footer(doc, report_type)

        # 保存文档
        file_path = self._save_document(doc, report_type, metadata)
        logger.info(f"Word 报告已生成: {file_path}")

        return file_path

    def _add_title(self, doc: Document, title: str):
        """添加标题"""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_metadata(self, doc: Document, metadata: Dict[str, Any]):
        """添加元数据"""
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if metadata.get("brand_name"):
            meta_paragraph.add_run(f"品牌/账号: {metadata['brand_name']}  ")

        if metadata.get("created_at"):
            meta_paragraph.add_run(f"生成时间: {metadata['created_at']}")
        else:
            meta_paragraph.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    def _add_summary(self, doc: Document, summary: str):
        """添加执行摘要"""
        heading = doc.add_heading("执行摘要", level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.runs[0].bold = True

    def _add_section(self, doc: Document, heading: str, content: str):
        """添加章节"""
        if heading:
            doc.add_heading(heading, level=1)

        if content:
            # 支持多段落内容
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

    def _add_footer(self, doc: Document, report_type: str):
        """添加页脚"""
        # 添加分隔线
        doc.add_paragraph()

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"—— 本报告由 AI 营销助手自动生成 ——\n"
            f"报告类型: {report_type}\n"
            f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _save_document(
        self, doc: Document, report_type: str, metadata: Optional[Dict[str, Any]]
    ) -> str:
        """保存文档"""
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        user_id = metadata.get("user_id", "unknown") if metadata else "unknown"
        filename = f"{report_type}_{user_id}_{timestamp}.docx"

        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)

        return file_path

    def get_download_url(self, file_path: str) -> str:
        """获取下载链接"""
        # 返回相对路径，前端会拼接成完整 URL
        return f"/data/reports/{os.path.basename(file_path)}"


class ReportContentBuilder:
    """报告内容构建器 - 根据不同报告类型构建内容"""

    @staticmethod
    def build_account_diagnosis_report(
        analysis: Dict[str, Any], user_id: str
    ) -> Dict[str, Any]:
        """构建账号诊断报告内容"""
        diagnosis = analysis.get("account_diagnosis", {})

        # 提取数据
        summary = diagnosis.get("summary", "暂无摘要")
        basic_info = diagnosis.get("basic_info", {})
        metrics = diagnosis.get("metrics", {})
        issues = diagnosis.get("issues", [])
        suggestions = diagnosis.get("suggestions", [])

        # 构建章节
        sections = []

        # 1. 账号概况
        sections.append({
            "heading": "一、账号概况",
            "content": f"账号名称: {basic_info.get('name', '未知')}\n"
            f"粉丝数量: {basic_info.get('fans', 0):,} 人\n"
            f"作品数量: {basic_info.get('works_count', 0)} 个\n"
            f"总播放量: {basic_info.get('total_views', 0):,} 次",
        })

        # 2. 核心指标
        sections.append({
            "heading": "二、核心指标分析",
            "content": f"平均播放量: {metrics.get('avg_views', 0):,} 次\n"
            f"平均点赞量: {metrics.get('avg_likes', 0):,} 个\n"
            f"平均评论量: {metrics.get('avg_comments', 0):,} 条\n"
            f"粉丝互动率: {metrics.get('like_rate', 0)}%\n"
            f"完播率: {metrics.get('completion_rate', 0)}%",
        })

        # 3. 问题诊断
        issues_content = ""
        if issues:
            for i, issue in enumerate(issues, 1):
                indicator = issue.get("indicator", "未知指标")
                msg = issue.get("msg", issue.get("value", ""))
                issues_content += f"{i}. {indicator}: {msg}\n"
        else:
            issues_content = "暂无明显问题"

        sections.append({
            "heading": "三、问题诊断",
            "content": issues_content,
        })

        # 4. 优化建议
        suggestions_content = ""
        if suggestions:
            for i, sug in enumerate(suggestions, 1):
                category = sug.get("category", "通用")
                content = sug.get("suggestion", "")
                suggestions_content += f"{i}. [{category}] {content}\n"
        else:
            suggestions_content = "暂无建议"

        sections.append({
            "heading": "四、优化建议",
            "content": suggestions_content,
        })

        return {
            "report_type": "account_diagnosis",
            "title": "账号诊断分析报告",
            "summary": summary,
            "sections": sections,
            "metadata": {
                "user_id": user_id,
                "brand_name": basic_info.get("name", ""),
                "report_type_display": "账号诊断报告",
            },
        }

    @staticmethod
    def build_marketing_strategy_report(
        analysis: Dict[str, Any], content: str, user_id: str
    ) -> Dict[str, Any]:
        """构建推广策略报告内容"""
        # 从 analysis 中提取策略信息
        angle = analysis.get("angle", "")
        reason = analysis.get("reason", "")

        sections = []

        # 1. 策略概述
        sections.append({
            "heading": "一、推广目标与定位",
            "content": analysis.get("target_audience", "根据目标人群特征进行精准推广"),
        })

        # 2. 内容策略
        sections.append({
            "heading": "二、内容策略",
            "content": f"推荐角度: {angle}\n\n策略说明: {reason}",
        })

        # 3. 渠道建议
        channels = analysis.get("recommended_channels", [])
        channels_content = "\n".join([f"- {ch}" for ch in channels]) if channels else "待分析"
        sections.append({
            "heading": "三、推广渠道建议",
            "content": channels_content,
        })

        # 4. 执行方案
        sections.append({
            "heading": "四、执行方案",
            "content": content[:2000] if content else "详见上方内容策略",
        })

        return {
            "report_type": "marketing_strategy",
            "title": "营销推广策略报告",
            "summary": f"为您的品牌/产品制定了详细的推广策略，包含{len(channels)}个推荐渠道",
            "sections": sections,
            "metadata": {
                "user_id": user_id,
                "brand_name": analysis.get("brand_name", ""),
                "report_type_display": "推广策略报告",
            },
        }

    @staticmethod
    def build_viral_prediction_report(
        analysis: Dict[str, Any], content: str, user_id: str
    ) -> Dict[str, Any]:
        """构建爆款预测报告内容"""
        prediction = analysis.get("viral_prediction", {})
        score = prediction.get("viral_score", 0)
        factors = prediction.get("viral_factors", [])

        sections = []

        # 1. 预测结论
        score_level = "高" if score >= 70 else "中" if score >= 40 else "低"
        sections.append({
            "heading": "一、爆款预测结论",
            "content": f"爆款潜力评分: {score}/100 ({score_level}潜力)\n"
            f"预测等级: {'🔥 高潜力' if score >= 70 else '📈 中等潜力' if score >= 40 else '📉 建议优化'}",
        })

        # 2. 爆款元素分析
        factors_content = ""
        if factors:
            for i, factor in enumerate(factors, 1):
                factors_content += f"{i}. {factor.get('element', '')}: {factor.get('analysis', '')}\n"
        else:
            factors_content = "暂无分析"

        sections.append({
            "heading": "二、爆款元素分析",
            "content": factors_content,
        })

        # 3. 优化建议
        suggestions = prediction.get("optimization_suggestions", [])
        suggestions_content = "\n".join([f"- {s}" for s in suggestions]) if suggestions else "暂无建议"

        sections.append({
            "heading": "三、优化建议",
            "content": suggestions_content,
        })

        return {
            "report_type": "viral_prediction",
            "title": "爆款预测分析报告",
            "summary": f"您的内容爆款潜力为 {score} 分，{score_level}潜力",
            "sections": sections,
            "metadata": {
                "user_id": user_id,
                "report_type_display": "爆款预测报告",
            },
        }

    @staticmethod
    def build_custom_report(
        title: str, content: str, analysis: Dict[str, Any], user_id: str
    ) -> Dict[str, Any]:
        """构建自定义报告内容"""
        sections = [
            {
                "heading": "一、内容详情",
                "content": content[:3000],
            },
            {
                "heading": "二、分析说明",
                "content": str(analysis)[:2000] if analysis else "无",
            },
        ]

        return {
            "report_type": "custom",
            "title": title or "内容分析报告",
            "summary": f"为您生成的内容报告 - {len(content)} 字符",
            "sections": sections,
            "metadata": {
                "user_id": user_id,
                "report_type_display": "内容报告",
            },
        }
